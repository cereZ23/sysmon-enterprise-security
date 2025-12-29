#!/usr/bin/env python3
"""
Markdown to DOCX Converter
Converts GUIDA-SYSMON-ENTERPRISE-IT.md to professional Word document.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def create_element(name):
    return OxmlElement(name)


def add_page_number(paragraph):
    """Add page number field to paragraph."""
    run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = create_element('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = create_element('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def setup_styles(doc):
    """Configure document styles."""
    styles = doc.styles

    # Title
    title = styles['Title']
    title.font.name = 'Calibri Light'
    title.font.size = Pt(26)
    title.font.color.rgb = RGBColor(0, 51, 102)
    title.font.bold = True

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = 'Calibri Light'
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(0, 51, 102)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(13)
    h2.font.color.rgb = RGBColor(0, 102, 153)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(11)
    h3.font.color.rgb = RGBColor(0, 102, 153)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)

    # Normal
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)

    # Create Code style
    if 'Code' not in [s.name for s in styles]:
        code = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = 'Consolas'
        code.font.size = Pt(8)
        code.paragraph_format.space_before = Pt(2)
        code.paragraph_format.space_after = Pt(2)
        code.paragraph_format.left_indent = Cm(0.3)


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def parse_table(lines):
    """Parse markdown table lines into rows."""
    rows = []
    for line in lines:
        if line.strip().startswith('|') and not re.match(r'^\|[\s\-:|]+\|$', line):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            rows.append(cells)
    return rows


def add_table(doc, rows):
    """Add formatted table."""
    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < len(table.rows[i].cells):
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                    if i == 0:
                        for run in p.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i == 0:
                    set_cell_shading(cell, '003366')

    doc.add_paragraph()


def add_code_block(doc, lines):
    """Add code block with background."""
    for line in lines:
        p = doc.add_paragraph(line, style='Code')
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F0F0F0')
        p._p.get_or_add_pPr().append(shading)


def process_inline(paragraph, text):
    """Process inline markdown (bold, italic, code)."""
    # Pattern to match **bold**, *italic*, `code`
    pattern = r'(\*\*.*?\*\*|`[^`]+`|\*[^*]+\*)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(153, 0, 0)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def convert_md_to_docx(md_path, docx_path):
    """Convert markdown to DOCX."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    setup_styles(doc)

    # Header
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.text = "Sysmon Enterprise Security Monitoring - Guida Tecnica"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # Footer with page numbers
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("Pagina ")
    add_page_number(fp)
    fp.add_run(" | Classificazione: Uso Interno")
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # Process content
    i = 0
    in_code = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_code_block(doc, code_lines)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            in_table = False
            rows = parse_table(table_lines)
            if rows:
                add_table(doc, rows)
            table_lines = []

        # Horizontal rule = page break
        if line.strip() == '---':
            doc.add_page_break()
            i += 1
            continue

        # Title (# )
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:].strip(), style='Title')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Heading 1 (## )
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue

        # Heading 2 (### )
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        # Heading 3 (#### )
        if line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        # Metadata (bold labels like **Version:**)
        if line.startswith('**') and ':**' in line:
            p = doc.add_paragraph()
            process_inline(p, line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Bullet list
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            process_inline(p, text)
            i += 1
            continue

        # Numbered list
        match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if match:
            text = match.group(2)
            p = doc.add_paragraph(style='List Number')
            process_inline(p, text)
            i += 1
            continue

        # Regular paragraph
        if line.strip():
            p = doc.add_paragraph()
            process_inline(p, line.strip())

        i += 1

    # Handle remaining table
    if in_table and table_lines:
        rows = parse_table(table_lines)
        if rows:
            add_table(doc, rows)

    doc.save(docx_path)
    print(f"✅ Documento generato: {docx_path}")


def main():
    script_dir = Path(__file__).parent
    md_file = script_dir / "GUIDA-SYSMON-ENTERPRISE-IT.md"
    docx_file = script_dir / "GUIDA-SYSMON-ENTERPRISE-IT.docx"

    if not md_file.exists():
        print(f"❌ File non trovato: {md_file}")
        return 1

    print(f"📄 Convertendo: {md_file.name}")
    convert_md_to_docx(str(md_file), str(docx_file))
    return 0


if __name__ == "__main__":
    exit(main())

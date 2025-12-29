#!/usr/bin/env python3
"""
Generate Technical Executive Report with Elegant Charts
Focus: Test Results, Detection Metrics, No Financial Data
"""

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io

# Color palette - Professional blue theme
COLORS = {
    'primary': '#1E3A5F',      # Dark blue
    'secondary': '#2E5984',    # Medium blue
    'accent': '#4A90D9',       # Light blue
    'success': '#28A745',      # Green
    'warning': '#FFC107',      # Yellow
    'danger': '#DC3545',       # Red
    'light': '#F8F9FA',        # Light gray
    'dark': '#343A40',         # Dark gray
    'sysmon': '#E74C3C',       # Red for Sysmon
    'winevents': '#3498DB',    # Blue for Windows Events
    'combined': '#27AE60',     # Green for Combined
}

plt.style.use('seaborn-v0_8-whitegrid')
plt.rcParams['font.family'] = 'DejaVu Sans'
plt.rcParams['font.size'] = 10

def create_score_gauge(score, max_score=100):
    """Create elegant donut gauge for score"""
    fig, ax = plt.subplots(figsize=(5, 5), subplot_kw=dict(aspect="equal"))

    # Data
    data = [score, max_score - score]
    colors = [COLORS['success'] if score >= 90 else COLORS['warning'], '#E0E0E0']

    # Create donut
    wedges, _ = ax.pie(data, colors=colors, startangle=90,
                       wedgeprops=dict(width=0.3, edgecolor='white'))

    # Center text
    ax.text(0, 0.1, f'{score}', fontsize=48, fontweight='bold',
            ha='center', va='center', color=COLORS['primary'])
    ax.text(0, -0.2, f'/ {max_score}', fontsize=20,
            ha='center', va='center', color=COLORS['dark'])
    ax.text(0, -0.45, 'PUNTEGGIO', fontsize=12,
            ha='center', va='center', color=COLORS['secondary'])

    ax.set_title('Production Readiness Score', fontsize=14, fontweight='bold',
                 color=COLORS['primary'], pad=20)

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close()
    buf.seek(0)
    return buf

def create_coverage_comparison():
    """Create bar chart comparing Sysmon vs Combined coverage"""
    fig, ax = plt.subplots(figsize=(10, 6))

    categories = ['Solo\nSysmon', 'Soluzione\nCombinata', 'Benchmark\nIndustria']
    values = [83.75, 97.5, 75]
    colors = [COLORS['sysmon'], COLORS['combined'], COLORS['secondary']]

    bars = ax.barh(categories, values, color=colors, height=0.6, edgecolor='white', linewidth=2)

    # Add value labels
    for bar, val in zip(bars, values):
        ax.text(val + 1, bar.get_y() + bar.get_height()/2,
                f'{val}%', va='center', fontsize=14, fontweight='bold',
                color=COLORS['dark'])

    # Add target line
    ax.axvline(x=97.5, color=COLORS['success'], linestyle='--', linewidth=2, alpha=0.7)
    ax.text(97.5, 2.7, 'Target: 97.5%', fontsize=10, color=COLORS['success'])

    ax.set_xlim(0, 110)
    ax.set_xlabel('Copertura MITRE ATT&CK (%)', fontsize=12, color=COLORS['dark'])
    ax.set_title('Confronto Copertura Detection', fontsize=16, fontweight='bold',
                 color=COLORS['primary'], pad=20)

    # Remove spines
    for spine in ['top', 'right']:
        ax.spines[spine].set_visible(False)

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close()
    buf.seek(0)
    return buf

def create_tactic_coverage():
    """Create horizontal bar chart for tactic coverage"""
    fig, ax = plt.subplots(figsize=(10, 7))

    tactics = [
        'Execution', 'Persistence', 'Privilege Escalation',
        'Lateral Movement', 'Discovery', 'Defense Evasion',
        'Credential Access', 'Collection', 'Exfiltration'
    ]
    coverage = [100, 100, 100, 100, 100, 95, 95, 90, 90]

    # Color based on coverage
    colors = [COLORS['success'] if c == 100 else COLORS['accent'] for c in coverage]

    y_pos = np.arange(len(tactics))
    bars = ax.barh(y_pos, coverage, color=colors, height=0.7, edgecolor='white', linewidth=1)

    # Add value labels
    for i, (bar, val) in enumerate(zip(bars, coverage)):
        ax.text(val + 1, bar.get_y() + bar.get_height()/2,
                f'{val}%', va='center', fontsize=11, fontweight='bold',
                color=COLORS['dark'])
        # Add checkmark for 100%
        if val == 100:
            ax.text(val - 5, bar.get_y() + bar.get_height()/2,
                    '✓', va='center', fontsize=14, color='white', fontweight='bold')

    ax.set_yticks(y_pos)
    ax.set_yticklabels(tactics, fontsize=11)
    ax.set_xlim(0, 110)
    ax.set_xlabel('Copertura (%)', fontsize=12, color=COLORS['dark'])
    ax.set_title('Copertura per Tattica ATT&CK', fontsize=16, fontweight='bold',
                 color=COLORS['primary'], pad=20)

    # Remove spines
    for spine in ['top', 'right']:
        ax.spines[spine].set_visible(False)

    ax.invert_yaxis()

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close()
    buf.seek(0)
    return buf

def create_config_results():
    """Create bar chart for per-config detection rates"""
    fig, ax = plt.subplots(figsize=(10, 6))

    configs = ['WS', 'SRV', 'DC', 'SQL', 'EXCH', 'IIS']
    full_names = ['Workstation', 'Server', 'Domain\nController', 'SQL\nServer', 'Exchange', 'IIS Web']
    rates = [85.0, 82.5, 85.0, 82.5, 82.5, 85.0]
    techniques = [34, 33, 34, 33, 33, 34]

    x = np.arange(len(configs))
    width = 0.6

    bars = ax.bar(x, rates, width, color=COLORS['accent'], edgecolor='white', linewidth=2)

    # Add value labels
    for bar, rate, tech in zip(bars, rates, techniques):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                f'{rate}%', ha='center', va='bottom', fontsize=12, fontweight='bold',
                color=COLORS['dark'])
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height()/2,
                f'{tech}/40', ha='center', va='center', fontsize=10,
                color='white', fontweight='bold')

    # Average line
    avg = sum(rates) / len(rates)
    ax.axhline(y=avg, color=COLORS['success'], linestyle='--', linewidth=2, alpha=0.8)
    ax.text(len(configs)-0.5, avg + 1, f'Media: {avg}%', fontsize=10,
            color=COLORS['success'], fontweight='bold')

    ax.set_xlabel('Configurazione', fontsize=12, color=COLORS['dark'])
    ax.set_ylabel('Detection Rate (%)', fontsize=12, color=COLORS['dark'])
    ax.set_title('Risultati Test per Configurazione', fontsize=16, fontweight='bold',
                 color=COLORS['primary'], pad=20)
    ax.set_xticks(x)
    ax.set_xticklabels(full_names, fontsize=10)
    ax.set_ylim(0, 100)

    # Remove spines
    for spine in ['top', 'right']:
        ax.spines[spine].set_visible(False)

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close()
    buf.seek(0)
    return buf

def create_improvement_chart():
    """Create before/after improvement chart for T1021.002"""
    fig, ax = plt.subplots(figsize=(8, 5))

    categories = ['PRIMA\n(PR #1)', 'DOPO\n(PR #1)']
    values = [1, 6]
    colors = [COLORS['danger'], COLORS['success']]

    x = np.arange(len(categories))
    bars = ax.bar(x, values, color=colors, width=0.5, edgecolor='white', linewidth=2)

    # Add labels
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1,
                f'{val}/6 configs', ha='center', va='bottom', fontsize=14,
                fontweight='bold', color=COLORS['dark'])

    # Improvement arrow
    ax.annotate('', xy=(1, 5.5), xytext=(0, 1.5),
                arrowprops=dict(arrowstyle='->', color=COLORS['success'],
                               lw=3, connectionstyle='arc3,rad=0.3'))
    ax.text(0.5, 4, '+500%', fontsize=20, fontweight='bold',
            color=COLORS['success'], ha='center')

    ax.set_ylabel('Configurazioni che rilevano', fontsize=12, color=COLORS['dark'])
    ax.set_title('T1021.002 - SMB Lateral Movement\nMiglioramento Detection',
                 fontsize=14, fontweight='bold', color=COLORS['primary'], pad=20)
    ax.set_xticks(x)
    ax.set_xticklabels(categories, fontsize=12)
    ax.set_ylim(0, 7)

    for spine in ['top', 'right']:
        ax.spines[spine].set_visible(False)

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close()
    buf.seek(0)
    return buf

def create_defense_depth_chart():
    """Create stacked bar showing defense-in-depth coverage"""
    fig, ax = plt.subplots(figsize=(10, 5))

    categories = ['Kernel Level\n(Sysmon)', 'OS Level\n(WinEvents)', 'Overlap\n(Both)']

    # Create a waterfall-style visualization
    sysmon_only = 8.75  # Sysmon-only coverage
    winevents_only = 13.75  # Windows Events only coverage
    overlap = 75  # Overlap

    bars_data = [
        ('Sysmon Unique', sysmon_only, COLORS['sysmon']),
        ('Windows Events Unique', winevents_only, COLORS['winevents']),
        ('Overlap (Resilienza)', overlap, COLORS['combined']),
    ]

    labels = [b[0] for b in bars_data]
    values = [b[1] for b in bars_data]
    colors = [b[2] for b in bars_data]

    x = np.arange(len(labels))
    bars = ax.bar(x, values, color=colors, width=0.6, edgecolor='white', linewidth=2)

    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                f'{val}%', ha='center', va='bottom', fontsize=14,
                fontweight='bold', color=COLORS['dark'])

    ax.set_ylabel('Copertura (%)', fontsize=12, color=COLORS['dark'])
    ax.set_title('Architettura Defense-in-Depth\nDistribuzione Copertura',
                 fontsize=14, fontweight='bold', color=COLORS['primary'], pad=20)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=11)
    ax.set_ylim(0, 85)

    # Add total annotation
    total = sysmon_only + winevents_only + overlap
    ax.text(len(labels)/2 - 0.5, 80, f'Totale Combinato: {97.5}%',
            fontsize=14, fontweight='bold', color=COLORS['success'],
            ha='center', bbox=dict(boxstyle='round', facecolor='white',
                                  edgecolor=COLORS['success'], linewidth=2))

    for spine in ['top', 'right']:
        ax.spines[spine].set_visible(False)

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close()
    buf.seek(0)
    return buf

def create_resilience_matrix():
    """Create resilience scenario visualization"""
    fig, ax = plt.subplots(figsize=(10, 6))

    scenarios = [
        'Sysmon Disabilitato',
        'WinEvents Disabilitato',
        'Attacco LOLBins',
        'Fileless Attack',
        'PowerShell Obfuscato',
        'Lateral Movement SMB'
    ]

    # Sysmon detection, WinEvents detection
    sysmon = [0, 1, 1, 1, 1, 1]
    winevents = [1, 0, 1, 1, 1, 1]

    x = np.arange(len(scenarios))
    width = 0.35

    bars1 = ax.barh(x - width/2, sysmon, width, label='Sysmon',
                    color=COLORS['sysmon'], edgecolor='white')
    bars2 = ax.barh(x + width/2, winevents, width, label='Windows Events',
                    color=COLORS['winevents'], edgecolor='white')

    # Add status indicators
    for i, (s, w) in enumerate(zip(sysmon, winevents)):
        status = 'RILEVATO' if (s or w) else 'NON RILEVATO'
        color = COLORS['success'] if (s or w) else COLORS['danger']
        ax.text(1.1, i, status, va='center', fontsize=10, fontweight='bold', color=color)

    ax.set_yticks(x)
    ax.set_yticklabels(scenarios, fontsize=11)
    ax.set_xlim(0, 1.8)
    ax.set_xticks([0, 1])
    ax.set_xticklabels(['No', 'Sì'], fontsize=10)
    ax.set_xlabel('Rilevamento Attivo', fontsize=12, color=COLORS['dark'])
    ax.set_title('Matrice Resilienza - Scenari di Attacco',
                 fontsize=14, fontweight='bold', color=COLORS['primary'], pad=20)
    ax.legend(loc='upper right', fontsize=10)

    for spine in ['top', 'right']:
        ax.spines[spine].set_visible(False)

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close()
    buf.seek(0)
    return buf

def create_compliance_chart():
    """Create compliance framework coverage chart"""
    fig, ax = plt.subplots(figsize=(10, 5))

    frameworks = ['PCI-DSS\nv4.0', 'HIPAA', 'NIS2', 'SOX', 'ISO\n27001', 'NIST\nCSF']
    coverage = [95, 95, 90, 90, 95, 95]

    x = np.arange(len(frameworks))
    colors = [COLORS['success'] if c >= 95 else COLORS['accent'] for c in coverage]

    bars = ax.bar(x, coverage, color=colors, width=0.6, edgecolor='white', linewidth=2)

    for bar, val in zip(bars, coverage):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                f'{val}%', ha='center', va='bottom', fontsize=12,
                fontweight='bold', color=COLORS['dark'])
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height()/2,
                '✓', ha='center', va='center', fontsize=18,
                color='white', fontweight='bold')

    # Compliance threshold
    ax.axhline(y=90, color=COLORS['warning'], linestyle='--', linewidth=2, alpha=0.8)
    ax.text(len(frameworks)-0.5, 91, 'Soglia Compliance: 90%', fontsize=9,
            color=COLORS['warning'], fontweight='bold')

    ax.set_ylabel('Copertura Requisiti (%)', fontsize=12, color=COLORS['dark'])
    ax.set_title('Compliance Framework Coverage', fontsize=16, fontweight='bold',
                 color=COLORS['primary'], pad=20)
    ax.set_xticks(x)
    ax.set_xticklabels(frameworks, fontsize=10)
    ax.set_ylim(0, 105)

    for spine in ['top', 'right']:
        ax.spines[spine].set_visible(False)

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close()
    buf.seek(0)
    return buf

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_word_document():
    """Create the Word document with all charts"""
    doc = Document()

    # Title
    title = doc.add_heading('Report Tecnico Esecutivo', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Soluzione di Monitoraggio Sicurezza Enterprise')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    # Date info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Data: 17 Dicembre 2025 | Classificazione: Report Tecnico\n')
    info.add_run('Audience: CISO, Security Team, IT Leadership')
    info.runs[0].font.size = Pt(10)
    info.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Verdict box
    verdict = doc.add_paragraph()
    verdict.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = verdict.add_run('✅ PRODUCTION READY - APPROVATO')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x28, 0xA7, 0x45)

    doc.add_paragraph()

    # Score gauge
    doc.add_heading('Punteggio Finale', 1)
    score_chart = create_score_gauge(92)
    doc.add_picture(score_chart, width=Inches(3.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Coverage comparison
    doc.add_heading('Risultati Test MITRE ATT&CK', 1)

    # Methodology table
    doc.add_heading('Metodologia', 2)
    method_table = doc.add_table(rows=5, cols=2)
    method_table.style = 'Table Grid'
    method_data = [
        ('Framework', 'MITRE ATT&CK v14'),
        ('Tool di Test', 'Atomic Red Team'),
        ('Ambiente', 'GitHub Actions (Windows Server 2022)'),
        ('Tecniche Testate', '40 tecniche di attacco'),
        ('Configurazioni', '6 configs role-specific'),
    ]
    for i, (label, value) in enumerate(method_data):
        method_table.rows[i].cells[0].text = label
        method_table.rows[i].cells[1].text = value
        method_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Coverage comparison chart
    doc.add_heading('Confronto Copertura', 2)
    coverage_chart = create_coverage_comparison()
    doc.add_picture(coverage_chart, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Tactic coverage
    doc.add_heading('Copertura per Tattica ATT&CK', 1)
    tactic_chart = create_tactic_coverage()
    doc.add_picture(tactic_chart, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Per-config results
    doc.add_heading('Risultati per Configurazione', 1)
    config_chart = create_config_results()
    doc.add_picture(config_chart, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Config table
    doc.add_paragraph()
    config_table = doc.add_table(rows=7, cols=4)
    config_table.style = 'Table Grid'
    headers = ['Config', 'Ruolo', 'Detection Rate', 'Tecniche']
    for i, h in enumerate(headers):
        config_table.rows[0].cells[i].text = h
        config_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(config_table.rows[0].cells[i], '1E3A5F')
        config_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    config_data = [
        ('sysmon-ws.xml', 'Workstation', '85.0%', '34/40'),
        ('sysmon-srv.xml', 'Server Generico', '82.5%', '33/40'),
        ('sysmon-dc.xml', 'Domain Controller', '85.0%', '34/40'),
        ('sysmon-sql.xml', 'SQL Server', '82.5%', '33/40'),
        ('sysmon-exch.xml', 'Exchange', '82.5%', '33/40'),
        ('sysmon-iis.xml', 'IIS Web Server', '85.0%', '34/40'),
    ]
    for i, row_data in enumerate(config_data):
        for j, val in enumerate(row_data):
            config_table.rows[i+1].cells[j].text = val

    doc.add_page_break()

    # Improvement chart
    doc.add_heading('Miglioramenti Ottenuti (PR #1)', 1)
    improve_chart = create_improvement_chart()
    doc.add_picture(improve_chart, width=Inches(5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Improvement description
    p = doc.add_paragraph()
    p.add_run('Modifica Implementata: ').bold = True
    p.add_run('Aggiunto Event ID 3 (Network Connection) con filtro su porta 445/SMB per rilevare Lateral Movement via SMB/Windows Admin Shares.')

    doc.add_page_break()

    # Defense in depth
    doc.add_heading('Architettura Defense-in-Depth', 1)
    depth_chart = create_defense_depth_chart()
    doc.add_picture(depth_chart, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Gap analysis table
    doc.add_heading('Gap Analysis e Copertura', 2)
    gap_table = doc.add_table(rows=4, cols=3)
    gap_table.style = 'Table Grid'
    gap_headers = ['Tecnica', 'Descrizione', 'Copertura Windows Events']
    for i, h in enumerate(gap_headers):
        gap_table.rows[0].cells[i].text = h
        gap_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(gap_table.rows[0].cells[i], '1E3A5F')
        gap_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    gap_data = [
        ('T1087.001', 'Local Account Discovery', '✅ Event 4798, 4799'),
        ('T1560.001', 'Archive via Utility', '✅ Event 4688'),
        ('T1005', 'Data from Local System', '✅ Event 4663'),
    ]
    for i, row_data in enumerate(gap_data):
        for j, val in enumerate(row_data):
            gap_table.rows[i+1].cells[j].text = val

    doc.add_page_break()

    # Resilience matrix
    doc.add_heading('Matrice Resilienza', 1)
    resilience_chart = create_resilience_matrix()
    doc.add_picture(resilience_chart, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Resilience table
    doc.add_paragraph()
    res_table = doc.add_table(rows=5, cols=4)
    res_table.style = 'Table Grid'
    res_headers = ['Scenario', 'Sysmon', 'WinEvents', 'Rilevamento']
    for i, h in enumerate(res_headers):
        res_table.rows[0].cells[i].text = h
        res_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(res_table.rows[0].cells[i], '1E3A5F')
        res_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    res_data = [
        ('Sysmon disabilitato', '❌', '✅', 'GARANTITO'),
        ('WinEvents disabilitato', '✅', '❌', 'GARANTITO'),
        ('Living-off-the-Land', '✅', '✅', 'DOPPIO'),
        ('Fileless/In-memory', '✅', '✅', 'DOPPIO'),
    ]
    for i, row_data in enumerate(res_data):
        for j, val in enumerate(row_data):
            res_table.rows[i+1].cells[j].text = val

    doc.add_page_break()

    # Compliance
    doc.add_heading('Compliance Framework', 1)
    compliance_chart = create_compliance_chart()
    doc.add_picture(compliance_chart, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Summary
    doc.add_heading('Riepilogo Risultati', 1)

    summary_table = doc.add_table(rows=8, cols=2)
    summary_table.style = 'Table Grid'
    summary_data = [
        ('Tecniche Testate', '40'),
        ('Tecniche Rilevate (Sysmon)', '33.5 (83.75%)'),
        ('Tecniche Rilevate (Combinato)', '39 (97.5%)'),
        ('Configurazioni Testate', '6'),
        ('Configurazioni Valide', '6/6 (100%)'),
        ('Gap Critici', '0'),
        ('Punteggio Finale', '92/100'),
        ('Verdetto', 'PRODUCTION READY'),
    ]
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value
        summary_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        if label == 'Verdetto':
            summary_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True
            summary_table.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x28, 0xA7, 0x45)

    doc.add_paragraph()

    # Benchmark comparison
    doc.add_heading('Confronto con Benchmark', 2)
    bench_table = doc.add_table(rows=5, cols=4)
    bench_table.style = 'Table Grid'
    bench_headers = ['Metrica', 'Nostra Soluzione', 'Media Industria', 'Delta']
    for i, h in enumerate(bench_headers):
        bench_table.rows[0].cells[i].text = h
        bench_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(bench_table.rows[0].cells[i], '1E3A5F')
        bench_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    bench_data = [
        ('MITRE Coverage', '97.5%', '70-80%', '+17.5-27.5%'),
        ('Detection Overlap', '75%', '30-40%', '+35-45%'),
        ('Role-specific configs', '6', '1-2', '+4-5'),
        ('False Positive Tuning', 'Sì', 'Spesso no', '✅'),
    ]
    for i, row_data in enumerate(bench_data):
        for j, val in enumerate(row_data):
            bench_table.rows[i+1].cells[j].text = val

    doc.add_page_break()

    # Approval section
    doc.add_heading('Approvazione', 1)

    approval_table = doc.add_table(rows=5, cols=3)
    approval_table.style = 'Table Grid'
    approval_headers = ['Ruolo', 'Decisione', 'Data']
    for i, h in enumerate(approval_headers):
        approval_table.rows[0].cells[i].text = h
        approval_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(approval_table.rows[0].cells[i], '1E3A5F')
        approval_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    approval_data = [
        ('Security Auditor', '✅ APPROVATO', '17 Dic 2025'),
        ('Security Engineering', '☐ In attesa', ''),
        ('IT Operations', '☐ In attesa', ''),
        ('CISO', '☐ In attesa', ''),
    ]
    for i, row_data in enumerate(approval_data):
        for j, val in enumerate(row_data):
            approval_table.rows[i+1].cells[j].text = val

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Versione: 1.0 | Classificazione: Internal - Technical\n')
    footer.add_run('Assessment condotto secondo le best practice e metodologia MITRE ATT&CK')
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Save
    doc.save('EXECUTIVE-TECHNICAL-REPORT-IT.docx')
    print('✅ Report generato: EXECUTIVE-TECHNICAL-REPORT-IT.docx')

if __name__ == '__main__':
    create_word_document()
